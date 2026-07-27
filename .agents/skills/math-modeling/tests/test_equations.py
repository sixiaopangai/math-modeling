import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from lxml import etree


SCRIPTS = Path(__file__).resolve().parents[1] / "tools" / "docx" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from equations import (
    OMML_NS,
    _replace_conversion_pair,
    _sha256,
    latex2omml,
    latex_to_docx,
    replace_placeholder,
    verify_conversion,
)


class EquationConversionTests(unittest.TestCase):
    def test_physics_greek_letters_relations_and_inverse_trig_are_supported(self):
        xml = etree.fromstring(
            latex2omml(r"\nu \approx \mu,\quad \theta=\arcsin(x)+\arccos(y)+\arctan(z)")
        )
        text = "".join(xml.itertext())
        for expected in ("ν", "≈", "μ", "arcsin", "arccos", "arctan"):
            self.assertIn(expected, text)

    def test_common_physics_constants_delimiters_and_number_sets_are_supported(self):
        xml = etree.fromstring(
            latex2omml(r"\hbar\omega,\quad \ell\in\mathbb{R},\quad \langle x\rangle=30^\circ")
        )
        text = "".join(xml.itertext())
        for expected in ("ℏ", "ω", "ℓ", "∈", "R", "⟨", "⟩", "°"):
            self.assertIn(expected, text)

    def test_common_commands_are_not_silently_corrupted(self):
        xml = etree.fromstring(latex2omml(r"x_1,\ldots,x_n"))
        text = "".join(xml.itertext())
        self.assertIn("…", text)
        self.assertNotIn("ldots", text)

    def test_nth_root_contains_visible_degree(self):
        xml = etree.fromstring(latex2omml(r"\sqrt[3]{x}"))
        degree = xml.find(f".//{{{OMML_NS}}}deg")
        self.assertIsNotNone(degree)
        self.assertIn("3", "".join(degree.itertext()))

    def test_matrix_generates_matrix_rows(self):
        xml = etree.fromstring(latex2omml(r"\begin{bmatrix}a&b\\c&d\end{bmatrix}"))
        rows = xml.findall(f".//{{{OMML_NS}}}mr")
        self.assertEqual(len(rows), 2)

    def test_unknown_command_fails_explicitly(self):
        with self.assertRaisesRegex(ValueError, "不支持"):
            latex2omml(r"\unknowncommand{x}")

    def test_unbalanced_group_fails_explicitly(self):
        with self.assertRaisesRegex(ValueError, "括号"):
            latex2omml(r"x_{i")
        with self.assertRaisesRegex(ValueError, "括号"):
            latex2omml("x}")

    def test_common_modeling_operators_and_styles_are_supported(self):
        xml = etree.fromstring(
            latex2omml(r"\min_x \operatorname{RMSE}(x)+\log(x)+\mathbf{w}^T x")
        )
        text = "".join(xml.itertext())
        for expected in ("min", "RMSE", "log", "w"):
            self.assertIn(expected, text)

    def test_cases_environment_generates_two_rows(self):
        xml = etree.fromstring(
            latex2omml(r"\begin{cases}x,&x>0\\0,&x\le 0\end{cases}")
        )
        rows = xml.findall(f".//{{{OMML_NS}}}mr")
        self.assertEqual(len(rows), 2)

    def test_replaces_every_matching_placeholder(self):
        doc = Document()
        doc.add_paragraph("第一处 EQ")
        doc.add_paragraph("第二处 EQ")

        replaced = replace_placeholder(doc, "EQ", "x")

        self.assertEqual(replaced, 2)
        self.assertFalse(any("EQ" in p.text for p in doc.paragraphs))

    def test_replaces_placeholder_inside_table_cell(self):
        doc = Document()
        cell = doc.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "指标 EQ"

        replaced = replace_placeholder(doc, "EQ", "x")

        self.assertEqual(replaced, 1)
        self.assertNotIn("EQ", cell.text)

    def test_converts_complete_latex_to_docx_with_pandoc_atomically(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "sections").mkdir()
            (root / "sections" / "body.tex").write_text(
                r"\section{模型}$x^2$", encoding="utf-8"
            )
            source = root / "main.tex"
            source.write_text(
                r"\documentclass{article}\begin{document}"
                "\\alpha% 注释\nbeta\n"
                r"\begin{verbatim}\input{ghost}\end{verbatim}"
                r"\input{sections/body}\end{document}",
                encoding="utf-8",
            )
            template = root / "template.docx"
            Document().save(template)
            output = root / "paper.docx"
            observed = []
            flattened_sources = []

            def fake_run(command, **kwargs):
                observed.append((command, kwargs))
                flattened_sources.append(
                    Path(command[command.index("--citeproc") + 1]).read_text(
                        encoding="utf-8"
                    )
                )
                generated = Path(command[command.index("--output") + 1])
                document = Document()
                document.add_paragraph("converted")
                document.save(generated)
                return subprocess.CompletedProcess(
                    command, 0, stdout="", stderr="conversion warning"
                )

            with patch("equations.shutil.which", return_value="pandoc"), patch(
                "equations.subprocess.run", side_effect=fake_run
            ):
                with self.assertRaisesRegex(RuntimeError, "警告"):
                    latex_to_docx(source, output, template)
                self.assertFalse(output.exists())
                result = latex_to_docx(
                    source,
                    output,
                    template,
                    timeout=45,
                    allow_warnings=[r"conversion warning"],
                    override_reason="确认该警告不影响公式转换",
                )

            command, options = observed[-1]
            self.assertEqual(command[command.index("--from") + 1], "latex")
            self.assertIn("--citeproc", command)
            self.assertIn(f"--resource-path={source.parent}", command)
            self.assertIn("--reference-doc", command)
            self.assertEqual(options["cwd"], source.parent)
            self.assertEqual(options["timeout"], 45)
            self.assertTrue(output.is_file())
            self.assertEqual(result["warnings"], ["conversion warning"])
            self.assertNotIn(r"\input{sections/body}", flattened_sources[-1])
            self.assertIn(r"\section{模型}", flattened_sources[-1])
            self.assertIn("\\alpha% 注释\nbeta", flattened_sources[-1])
            self.assertIn(
                r"\begin{verbatim}\input{ghost}\end{verbatim}",
                flattened_sources[-1],
            )
            manifest = json.loads(
                output.with_suffix(".conversion.json").read_text(encoding="utf-8")
            )
            self.assertEqual(
                manifest["source_files"], ["main.tex", "sections/body.tex"]
            )
            self.assertEqual(
                manifest["warning_override"]["reason"],
                "确认该警告不影响公式转换",
            )
            verified = verify_conversion(output)
            self.assertTrue(verified["passed"], verified["issues"])

            manifest.pop("template")
            manifest.pop("template_sha256")
            output.with_suffix(".conversion.json").write_text(
                json.dumps(manifest, ensure_ascii=False), encoding="utf-8"
            )
            incomplete = verify_conversion(output)
            self.assertFalse(incomplete["passed"])
            self.assertTrue(
                any("缺少必填字段" in issue for issue in incomplete["issues"])
            )
            output.with_suffix(".conversion.json").write_text(
                json.dumps(
                    {
                        **manifest,
                        "template": str(template),
                        "template_sha256": _sha256(template),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            complete_manifest = json.loads(
                output.with_suffix(".conversion.json").read_text(encoding="utf-8")
            )
            for field, invalid in {
                "template": 123,
                "created_at": 123,
                "returncode": "0",
                "duration_seconds": "x",
                "reproduce": [],
            }.items():
                invalid_manifest = {**complete_manifest, field: invalid}
                output.with_suffix(".conversion.json").write_text(
                    json.dumps(invalid_manifest, ensure_ascii=False),
                    encoding="utf-8",
                )
                invalid_report = verify_conversion(output)
                self.assertFalse(
                    invalid_report["passed"],
                    f"{field}={invalid!r} 未被拒绝",
                )
            output.with_suffix(".conversion.json").write_text(
                json.dumps(complete_manifest, ensure_ascii=False),
                encoding="utf-8",
            )

            (root / "sections" / "body.tex").write_text(
                r"\section{模型}$changed$", encoding="utf-8"
            )
            stale = verify_conversion(output)
            self.assertFalse(stale["passed"])
            self.assertTrue(any("哈希" in issue for issue in stale["issues"]))

            original_docx = output.read_bytes()
            output.write_bytes(original_docx + b"tampered")
            tampered = verify_conversion(output)
            self.assertTrue(any("DOCX 哈希" in issue for issue in tampered["issues"]))
            self.assertEqual(
                {path.name for path in root.iterdir()},
                {
                    "main.tex",
                    "sections",
                    "template.docx",
                    "paper.docx",
                    "paper.conversion.json",
                },
            )

            with self.assertRaises(FileExistsError):
                latex_to_docx(source, output, template)

    def test_conversion_pair_restores_old_files_when_manifest_replace_fails(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "new.docx"
            Document().save(source)
            target = root / "paper.docx"
            Document().save(target)
            old_docx = target.read_bytes()
            manifest = root / "paper.conversion.json"
            manifest.write_text('{"old": true}', encoding="utf-8")
            real_replace = os.replace

            def fail_manifest(source_path, target_path):
                if (
                    Path(target_path) == manifest
                    and ".new-" in Path(source_path).name
                ):
                    raise OSError("simulated manifest failure")
                return real_replace(source_path, target_path)

            with patch("equations.os.replace", side_effect=fail_manifest):
                with self.assertRaisesRegex(OSError, "simulated"):
                    _replace_conversion_pair(
                        source, target, {"new": True}, overwrite=True
                    )

            self.assertEqual(target.read_bytes(), old_docx)
            self.assertEqual(manifest.read_text(encoding="utf-8"), '{"old": true}')


if __name__ == "__main__":
    unittest.main()
