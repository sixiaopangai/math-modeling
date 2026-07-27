import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "tools" / "docx" / "scripts"
sys.path.insert(0, str(SCRIPTS))

import paper_format as pf


class PaperFormatTests(unittest.TestCase):
    def _front_matter(self):
        doc = pf.new_document(contest="cumcm")
        pf.title(doc, "题目")
        pf.abstract_title(doc)
        pf.body(doc, "摘要正文")
        pf.keywords(doc, "优化；预测")
        return doc

    def test_reference_template_styles_are_kept_but_sample_body_is_cleared(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            source = Document()
            source.add_paragraph("模板示例正文，不应进入论文")
            source.save(template)

            doc = pf.new_document(contest="cumcm", template_path=template)

        self.assertNotIn("模板示例正文", "\n".join(p.text for p in doc.paragraphs))

    def test_official_fixed_template_content_can_be_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "official.docx"
            source = Document()
            source.add_paragraph("官方固定摘要页")
            source.save(template)

            doc = pf.new_document(
                contest="cumcm",
                template_path=template,
                preserve_template_content=True,
            )

        self.assertIn("官方固定摘要页", "\n".join(p.text for p in doc.paragraphs))

    def test_cumcm_structure_validator_requires_abstract_and_keywords(self):
        doc = pf.new_document(contest="cumcm")
        pf.title(doc, "题目")

        errors = pf.validate_paper_structure(doc, contest="cumcm")

        self.assertTrue(any("摘要" in error for error in errors))
        self.assertTrue(any("关键词" in error for error in errors))

    def test_complete_cumcm_front_matter_passes(self):
        doc = self._front_matter()

        errors = pf.validate_paper_structure(doc, contest="cumcm", quality_checks=False)

        self.assertEqual(errors, [])

    def test_quality_validation_reports_length_formula_figure_table_and_page_gaps(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(doc, contest="cumcm")

        for expected in ("15000", "公式", "图", "表", "渲染页数"):
            self.assertTrue(any(expected in issue for issue in issues), expected)
        self.assertTrue(any("低于质量目标 8" in issue for issue in issues))

    def test_other_contests_share_the_eight_figure_default(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="mcm-icm",
            min_content_units=0,
            min_equations=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("低于质量目标 8" in issue for issue in issues))

    def test_table_caption_must_be_referenced_in_body(self):
        doc = self._front_matter()
        pf.body(doc, "正文没有引用下面的表格。")
        pf.three_line_table(doc, [["变量", "值"], ["x", "1"]])
        doc.add_paragraph("表1 参数结果")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("表1" in issue and "正文" in issue for issue in issues))

    def test_reference_list_and_body_citations_are_bidirectionally_checked(self):
        doc = self._front_matter()
        pf.body(doc, "已有研究支持该方法[1]，但错误引用了[3]。")
        doc.add_paragraph("参考文献")
        doc.add_paragraph("[1] A. Author. A useful paper.")
        doc.add_paragraph("[2] B. Author. An uncited paper.")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertTrue(any("[3]" in issue and "参考文献表" in issue for issue in issues))
        self.assertTrue(any("[2]" in issue and "未在正文引用" in issue for issue in issues))

    def test_compound_reference_citations_are_recognized(self):
        doc = self._front_matter()
        pf.body(doc, "相关方法见文献[1, 2]及文献[3-4]。")
        doc.add_paragraph("参考文献")
        for number in range(1, 5):
            doc.add_paragraph(f"[{number}] Reference {number}.")

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            require_rendered_pages=False,
        )

        self.assertFalse(any("未在正文引用" in issue for issue in issues))

    def test_rendered_page_limits_distinguish_target_from_official_maximum(self):
        doc = self._front_matter()

        issues = pf.validate_paper_structure(
            doc,
            contest="cumcm",
            min_content_units=0,
            min_equations=0,
            min_figures=0,
            min_tables=0,
            rendered_pages=31,
        )

        self.assertTrue(any("官方上限" in issue and "30" in issue for issue in issues))

    def test_safe_save_rejects_skill_root(self):
        doc = self._front_matter()

        with self.assertRaisesRegex(ValueError, "PROJECT_ROOT"):
            pf.save_document(doc, pf.SKILL_ROOT, contest="cumcm")

    def test_completion_gate_rejects_incomplete_docx(self):
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "完整论文.docx"
            self._front_matter().save(path)

            report = pf.validate_document(path, contest="cumcm", rendered_pages=7)

        self.assertFalse(report["passed"])
        self.assertLess(report["metrics"]["content_units"], 15000)
        self.assertTrue(any("15000" in issue for issue in report["issues"]))


if __name__ == "__main__":
    unittest.main()
